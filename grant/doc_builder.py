"""
Builds a clean Word (.docx) document from the generated grant text.
"""
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_title(doc: Document, grant_name: str):
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"NEXORA GRANT APPLICATION")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(grant_name.upper())
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    doc.add_paragraph()


def _add_divider(doc: Document):
    para = doc.add_paragraph("─" * 50)
    para.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    para.runs[0].font.size = Pt(9)


def _add_question_block(doc: Document, question: str, answer: str):
    _add_divider(doc)

    q_para = doc.add_paragraph()
    q_run = q_para.add_run(question.strip())
    q_run.bold = True
    q_run.font.size = Pt(12)
    q_run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    doc.add_paragraph()

    for line in answer.strip().split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph(line)
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()


def build_word_document(grant_name: str, raw_text: str) -> io.BytesIO:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    _add_title(doc, grant_name)

    # Parse sections: split on "### " or numbered "1. " headings
    blocks = _parse_sections(raw_text)

    if blocks:
        for question, answer in blocks:
            _add_question_block(doc, question, answer)
    else:
        # Fallback: dump raw text
        for line in raw_text.split("\n"):
            doc.add_paragraph(line.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _parse_sections(text: str) -> list:
    """
    Extract (question, answer) pairs from the generated text.
    Handles '### Question' format and plain numbered questions.
    """
    results = []

    # Try ### headers first
    pattern = r"###\s+(.+?)(?=\n###|\Z)"
    matches = re.findall(pattern, text, re.DOTALL)

    if matches:
        for match in matches:
            lines = match.strip().split("\n", 1)
            question = lines[0].strip()
            answer = lines[1].strip() if len(lines) > 1 else ""
            if question:
                results.append((question, answer))
        return results

    # Fallback: split on blank lines — treat short lines as section headers
    blocks = re.split(r"\n{2,}", text.strip())
    for block in blocks:
        lines = block.strip().split("\n")
        if len(lines) >= 2 and len(lines[0]) < 120:
            results.append((lines[0].strip(), "\n".join(lines[1:]).strip()))
        elif lines:
            results.append(("", block.strip()))

    return results
